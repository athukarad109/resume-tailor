"""Generate .docx documents (e.g. cover letter) from plain text."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt


def _add_paragraph_with_bold(doc: Document, text: str) -> None:
    """Add a paragraph, rendering **bold** as bold runs."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    # Split by ** to alternate plain and bold
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def cover_letter_text_to_docx_bytes(cover_letter_text: str) -> bytes:
    """Convert cover letter plain text to a .docx file (paragraphs; **bold** supported)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        _add_paragraph_with_bold(doc, para_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
